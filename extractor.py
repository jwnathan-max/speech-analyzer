"""
extractor.py
연설 파일(오디오/PDF/TXT)에서 텍스트를 추출하고 정제하는 모듈.

- 오디오 (mp3, m4a, wav, ogg, flac): OpenAI Whisper API → STT
- PDF (일반): pdfplumber → 텍스트 정제
- PDF (JW 커스텀폰트): pypdfium2로 이미지 렌더링 → Claude Vision으로 OCR
- TXT: 직접 읽기

JW PDF 감지:
  - "No.X-KO" 패턴이 존재하면 S-34 골자 모음집으로 판단
  - 텍스트가 깨진(garbled) 경우에도 Vision fallback 처리
"""

import base64
import io
import os
import re
from pathlib import Path
from typing import Optional

import anthropic
import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from openai import OpenAI

# 지원 파일 확장자
AUDIO_EXTENSIONS = {".mp3", ".m4a", ".wav", ".ogg", ".flac", ".webm"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
TEXT_EXTENSIONS = {".txt"}

# Whisper API 파일 크기 제한 (25MB)
WHISPER_SIZE_LIMIT = 25 * 1024 * 1024


# ── Public API ─────────────────────────────────────────────────────────────


def get_text_from_file(file_path: str, page_index: Optional[int] = None) -> str:
    """
    파일 경로를 받아 텍스트를 추출하고 정제된 문자열을 반환한다.
    page_index: JW 골자 PDF에서 특정 페이지만 추출할 때 사용.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

    ext = path.suffix.lower()

    if ext in AUDIO_EXTENSIONS:
        return _extract_from_audio(path)
    elif ext in PDF_EXTENSIONS:
        return _extract_from_pdf(path, page_index=page_index)
    elif ext in DOCX_EXTENSIONS:
        return _extract_from_docx(path)
    elif ext in TEXT_EXTENSIONS:
        return _extract_from_txt(path)
    else:
        raise ValueError(
            f"지원하지 않는 파일 형식입니다: {ext}\n"
            f"지원 형식: {AUDIO_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | TEXT_EXTENSIONS}"
        )


def list_jw_outlines(pdf_path: str) -> dict[int, int]:
    """
    JW 골자 모음집(S-34 등)에서 골자 번호 → 페이지 인덱스 맵을 반환한다.
    "No.X-KO" 패턴이 없으면 빈 dict 반환 (일반 PDF로 취급).
    반환 예: {1: 0, 2: 2, 3: 4, ...}
    """
    outline_map: dict[int, int] = {}
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            raw = page.extract_text() or ""
            m = re.search(r"No\.(\d+)-KO", raw)
            if m:
                outline_map[int(m.group(1))] = i
    return outline_map


# ── 오디오 ─────────────────────────────────────────────────────────────────


def _extract_from_audio(path: Path) -> str:
    """OpenAI Whisper API를 사용해 오디오를 텍스트로 변환한다."""
    file_size = path.stat().st_size
    if file_size > WHISPER_SIZE_LIMIT:
        raise ValueError(
            f"오디오 파일이 너무 큽니다 ({file_size / 1024 / 1024:.1f}MB). "
            f"Whisper API 제한: 25MB 이하."
        )

    client = OpenAI()

    with open(path, "rb") as audio_file:
        transcript = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            language="ko",
            response_format="text",
        )

    return _clean_text(str(transcript))


# ── PDF ────────────────────────────────────────────────────────────────────


def _extract_from_pdf(path: Path, page_index: Optional[int] = None) -> str:
    """
    PDF에서 텍스트를 추출한다.
    page_index 지정 시 해당 페이지만 추출 (JW 골자 단일 페이지용).
    텍스트가 깨진 경우 Claude Vision으로 자동 fallback.
    """
    if page_index is not None:
        # 특정 페이지만 Vision으로 추출 (JW 골자)
        return _extract_with_vision(str(path), [page_index])

    # 일반 PDF: 먼저 pdfplumber로 시도
    pages_text = []
    with pdfplumber.open(str(path)) as pdf:
        page_list = pdf.pages
        for page in page_list:
            height = page.height
            cropped = page.within_bbox(
                (0, height * 0.10, page.width, height * 0.90)
            )
            text = cropped.extract_text(x_tolerance=3, y_tolerance=3)
            if text:
                pages_text.append(text)

    raw = "\n".join(pages_text)

    # 텍스트 깨짐 감지 → Vision fallback
    if _is_garbled(raw):
        with pdfplumber.open(str(path)) as pdf:
            all_indices = list(range(len(pdf.pages)))
        return _extract_with_vision(str(path), all_indices)

    return _clean_text(raw)


def _is_garbled(text: str) -> bool:
    """
    텍스트가 JW 커스텀폰트 인코딩으로 깨진 경우를 감지한다.
    Latin-1 보충(U+0080~U+00FF) 문자가 전체의 15% 이상이면 깨진 것으로 판단.
    """
    if not text:
        return False
    non_ascii = [c for c in text if ord(c) > 0x7F]
    if not non_ascii:
        return False
    latin1_supplement = sum(1 for c in non_ascii if 0x0080 <= ord(c) <= 0x00FF)
    return (latin1_supplement / len(non_ascii)) > 0.15


# ── Claude Vision OCR ──────────────────────────────────────────────────────


def _render_page_to_base64(pdf_path: str, page_index: int) -> str:
    """pypdfium2로 PDF 페이지를 렌더링해 base64 PNG로 반환한다."""
    doc = pdfium.PdfDocument(pdf_path)
    page = doc[page_index]
    # 150 DPI 수준 (A4 기준 약 1240×1754px)
    bitmap = page.render(scale=1.75)
    pil_image = bitmap.to_pil()
    doc.close()

    buf = io.BytesIO()
    pil_image.save(buf, format="PNG")
    return base64.standard_b64encode(buf.getvalue()).decode("utf-8")


def _extract_with_vision(pdf_path: str, page_indices: list[int]) -> str:
    """
    Claude Vision API를 사용해 PDF 페이지 이미지에서 텍스트를 추출한다.
    여러 페이지는 순서대로 이어 붙인다.
    """
    client = anthropic.Anthropic()
    all_text: list[str] = []

    for idx in page_indices:
        img_b64 = _render_page_to_base64(pdf_path, idx)

        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": "image/png",
                                "data": img_b64,
                            },
                        },
                        {
                            "type": "text",
                            "text": (
                                "이 이미지는 여호와의 증인 한국어 연설 골자(강연 원고)입니다. "
                                "이미지에 보이는 모든 텍스트를 정확히 그대로 추출해 주세요. "
                                "성구 참조(예: 마태 5:3, 시편 23:1)도 빠짐없이 포함하세요. "
                                "레이아웃(들여쓰기, 줄바꿈)을 최대한 보존하세요. "
                                "페이지 하단의 'No.X-KO' 코드와 저작권 줄은 제외하세요. "
                                "추출한 텍스트만 출력하고 다른 설명은 하지 마세요."
                            ),
                        },
                    ],
                }
            ],
        )
        all_text.append(message.content[0].text.strip())

    return _clean_text("\n\n".join(all_text))


# ── DOCX ───────────────────────────────────────────────────────────────────


def _extract_from_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return _clean_text("\n".join(parts))


# ── TXT ────────────────────────────────────────────────────────────────────


def _extract_from_txt(path: Path) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    return _clean_text(raw)


# ── 텍스트 정제 ────────────────────────────────────────────────────────────


def _clean_text(text: str) -> str:
    text = re.sub(r"-\n", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(lines).strip()


# ── 파일 목록 유틸 ─────────────────────────────────────────────────────────


def list_input_files(folder: str = "./input_files") -> list[dict]:
    folder_path = Path(folder)
    if not folder_path.exists():
        return []

    supported = AUDIO_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | TEXT_EXTENSIONS
    files = []
    for f in sorted(folder_path.iterdir()):
        if f.is_file() and f.suffix.lower() in supported:
            ext = f.suffix.lower()
            if ext in AUDIO_EXTENSIONS:
                ftype = "오디오"
            elif ext in PDF_EXTENSIONS:
                ftype = "PDF"
            elif ext in DOCX_EXTENSIONS:
                ftype = "워드"
            else:
                ftype = "텍스트"
            files.append({"name": f.name, "path": str(f), "type": ftype})
    return files
